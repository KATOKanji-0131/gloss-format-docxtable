import docx
from docx.shared import Pt
from docx.shared import Mm
import math
import os.path
import pandas as pd
import re
import varibles

with open("example.txt") as f:
    text_lines = f.readlines()

# グロスと略語のリストを作り、略語の文字数でソート。文字数の長い順に処理しないと、NONPASTのPASTに"PAST"がヒットしてしまうため。
abbreviations = pd.read_table("gloss_abbreviations.tsv", header=None).values.tolist()
abbreviations = sorted(abbreviations, reverse=True, key=lambda x: len(x[0]))
abbreviations_used = []

morpheme_texts = []
gloss_texts = []
translation_texts = []

morph_spcf = varibles.morph_spcf # 形態素行の指定子
gl_spcf = varibles.gl_spcf # グロス行の指定子
trsl_spcf = varibles.trsl_spcf # 訳行の指定子

# テキストから行を抽出してリストに格納する
for line in text_lines:
    if re.match(morph_spcf, line) is not None:
        line = re.sub(f"{morph_spcf}\s*", r"", line.rstrip("\n"))
        morpheme_texts.append(line.split())
    if re.match(gl_spcf, line) is not None:
        line = re.sub(f"{gl_spcf}\s*", r"", line.rstrip("\n"))
        gloss_texts.append(line.split())
    if re.match(trsl_spcf, line) is not None:
        line = re.sub(f"{trsl_spcf}\s*", r"", line.rstrip("\n"))
        translation_texts.append(line)

# ワードドキュメント作成
doc = docx.Document()

# 表整形の処理。各ループは例文一つに相当。
exno = 0 # 例文番号
for mp_block, gl_block, tr in zip(morpheme_texts, gloss_texts, translation_texts):
    table = doc.add_table(rows=2, cols=9)
    # table.autofit = False
    table.allow_autofit = False
    l = 0  # 例文の行数。はみ出しそうになったらlを+1して、次の行に続ける。

    # 一番左上のセルには例文番号を入れる
    table.rows[0].cells[0].text = "(" + str(exno+1) + ")"
    table.rows[0].cells[0].width = Mm(10)

    m = 0 # 形態素の番号。
    cell_width_sum = 0 # セルの長さの合計値。これがページ幅を超える時に改行する。

    for mp, gl in zip(mp_block, gl_block):
        # 表幅がページ幅を超えた時の処理
        if m > 7 or cell_width_sum > Mm(220):
            row = table.add_row()
            row = table.add_row()
            m = 0
            l += 1
            cell_width_sum = 0
        col_length = max(len(mp), len(gl))

        # 行数*2 番目の行に形態素、行数*2-1 番目の行にグロスをいれる。
        mp_row = table.rows[l * 2]
        gl_row = table.rows[l * 2 + 1]
        m += 1
        mp_row.cells[m].text = mp

        # もしglがグロス一覧にあるなら、その部分をスモールキャップに変換して、セルに追加する。
        for abb in abbreviations:
            if abb[0] not in gl:
                continue
            # 略号でない部分 (=prefixとsuffixの整形)
            try:
                prefix = re.match(f".+?(?={abb[0]})", gl).group()
            except AttributeError:
                prefix = ""
            try:
                suffix = re.match(f"(?<={abb[0]}).+?$", gl).group()
            except AttributeError:
                suffix = ""
            para = gl_row.cells[m].paragraphs[0]
            para.add_run(prefix)
            gloss = abb[0]
            p = para.add_run(gloss.lower())
            p.font.small_caps = True
            para.add_run(suffix)

            # 使用したabbreviationをリストに入れておき、後で略号一覧を出力する
            if abb not in abbreviations_used:
                abbreviations_used.append(abb)
            # print(abbreviations_used)

            # グロスが一つでもヒットしたら処理を終了したいのでbreak
            break

        # グロスに略号が含まれない場合、グロスをそのままセルに入れる
        if gl_row.cells[m].text == "":
            gl_row.cells[m].text = gl

        # 形態素とグロスのうち長い方に合わせてセル幅を設定する
        # ToDo: 一定の値を下回って小さくすることができなさそう？
        col_length = max(len(mp_row.cells[m].text), len(gl_row.cells[m].text))

        mp_row.cells[m].width = Pt(10*col_length)
        gl_row.cells[m].width = Pt(10*col_length)

        cell_width_sum += 10*col_length

    row = table.add_row()
    table.rows[-1].cells[1].text = tr
    # ToDo: 列高さ指定できるように
    for row in table.rows:
        row.height = Pt(8)
    doc.add_paragraph("")
    exno += 1

# 略号一覧の整形
abbreviations_used = sorted(abbreviations_used, key=lambda x: x[0])
abb_row_length = math.ceil(len(abbreviations_used)/2)
abb_table = doc.add_table(rows=abb_row_length, cols=4)
abb_table.allow_autofit = False
for row, i in zip(abb_table.rows, range(abb_row_length)):
    row.cells[0].text = abbreviations_used[i][0]
    row.cells[1].text = abbreviations_used[i][1]
    row.cells[2].text = abbreviations_used[i+abb_row_length][0]
    row.cells[3].text = abbreviations_used[i+abb_row_length][1]
abb_table.columns[0].width = Pt(30)

# ToDo セル幅を個別に指定できるように
# ToDo 余白を指定できるように
# ToDo 数字の後ろに英字を出力できるように

# ファイル名を指定して保存
file_no = 1
if os.path.exists(f"{varibles.file_name}.docx"):
    while os.path.exists(f"{varibles.file_name}({file_no}).docx") is True:
        file_no += 1
    doc.save(f"{varibles.file_name}({file_no}).docx")
else:
    doc.save(f"{varibles.file_name}.docx")

# doc.save(f"{varibles.file_name}.docx")
